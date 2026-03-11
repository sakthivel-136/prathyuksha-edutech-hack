import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_documentation():
    doc = Document()

    # Font helper
    def set_font(run, size=11, bold=False):
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = bold

    # Title
    title = doc.add_heading('', 0)
    run = title.add_run("VantageEdu: Integrated Academic & Examination Management System")
    set_font(run, size=20, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------
    # 1. ABSTRACT
    # ------------------------------------------------
    doc.add_heading("1. Abstract", level=1)

    abstract_text = (
        "VantageEdu is a modern academic and examination management platform designed to simplify "
        "the everyday administrative operations of educational institutions. Many colleges still rely "
        "on spreadsheets, manual registers, and disconnected software tools to manage student records, "
        "exam schedules, seating arrangements, and hall ticket distribution. These traditional methods "
        "often lead to data inconsistencies, human errors, and unnecessary delays.\n\n"

        "The purpose of VantageEdu is to centralize all academic and examination related processes into "
        "a single digital platform. The system allows administrators, faculty members, and students to "
        "interact with the platform through role-based dashboards. Administrators can manage academic "
        "calendars, examination schedules, and seating allocations, while students can easily access "
        "their exam details, download hall tickets, and stay informed about academic events.\n\n"

        "By automating repetitive administrative tasks and providing real-time data access, VantageEdu "
        "improves efficiency, reduces manual workload, and increases transparency within educational "
        "institutions. The platform is designed to be scalable, reliable, and suitable for deployment "
        "in colleges and universities where managing large volumes of student data is essential."
    )

    p = doc.add_paragraph()
    run = p.add_run(abstract_text)
    set_font(run)

    # ------------------------------------------------
    # 2. PROBLEM STATEMENT
    # ------------------------------------------------
    doc.add_heading("2. Problem Statement", level=1)

    problem_text = (
        "Educational institutions manage a wide range of academic activities including student records, "
        "attendance tracking, examination scheduling, seating arrangements, and hall ticket distribution. "
        "In many institutions these processes are still handled manually or using multiple disconnected "
        "software systems. As a result, staff members spend a significant amount of time entering data, "
        "verifying information, and correcting errors.\n\n"

        "One of the most time-consuming tasks is preparing examination seating arrangements. For large "
        "student populations this process can take several days of manual planning and still result in "
        "mistakes such as duplicate seat numbers or incorrect room allocations. Similarly, generating "
        "and distributing hall tickets manually creates unnecessary workload for administrative staff.\n\n"

        "Another major challenge is the absence of intelligent monitoring systems that can identify "
        "students who may be academically struggling. Without proper data analysis tools, students at "
        "risk of failing are often identified too late for effective academic intervention.\n\n"

        "These issues highlight the need for an integrated digital solution that can automate academic "
        "administration processes and improve overall efficiency."
    )

    p = doc.add_paragraph()
    run = p.add_run(problem_text)
    set_font(run)

    # ------------------------------------------------
    # 3. PROPOSED SOLUTION
    # ------------------------------------------------
    doc.add_heading("3. Proposed Solution", level=1)

    solution_text = (
        "VantageEdu addresses these challenges by introducing a centralized academic management system "
        "that integrates multiple administrative processes into a single platform. The system allows "
        "administrators to manage student records, academic calendars, examination schedules, and "
        "seating allocations through an intuitive interface.\n\n"

        "The platform automatically generates seating arrangements based on the available classroom "
        "capacity and the number of students registered for each subject. This reduces manual effort "
        "and ensures fair distribution of students across examination halls.\n\n"

        "Students can access their personal dashboards where they can view examination schedules, "
        "download hall tickets, and receive updates regarding academic events. The system also provides "
        "analytical tools that monitor student performance and identify individuals who may require "
        "additional academic support.\n\n"

        "Through automation, intelligent analytics, and centralized data management, VantageEdu "
        "transforms traditional academic administration into an efficient digital workflow."
    )

    p = doc.add_paragraph()
    run = p.add_run(solution_text)
    set_font(run)

    # ------------------------------------------------
    # 4. SYSTEM ARCHITECTURE
    # ------------------------------------------------
    doc.add_heading("4. System Architecture", level=1)

    arch_text = (
        "The architecture of VantageEdu follows a modular design in which multiple functional "
        "components interact with a centralized database. Each module performs a specific task "
        "such as student management, examination scheduling, seating allocation, and hall ticket "
        "generation.\n\n"

        "The system architecture consists of three primary layers:\n"
        "1. User Interface Layer – Provides dashboards for students, administrators, and faculty.\n"
        "2. Application Processing Layer – Handles business logic, scheduling algorithms, and system workflows.\n"
        "3. Database Layer – Stores student records, exam data, and system logs.\n\n"

        "This layered architecture ensures system scalability, maintainability, and reliable "
        "data management across all academic processes."
    )

    p = doc.add_paragraph()
    run = p.add_run(arch_text)
    set_font(run)

    # Use the flowchart image
    image_path = "/Users/rekha/.gemini/antigravity/brain/090f77a0-a0bf-488a-b422-56d169e63ad7/vantageedu_flowchart_architecture_1772986012847.png"
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("System Architecture Diagram (Placeholder)")

    # ------------------------------------------------
    # 5. MODULES
    # ------------------------------------------------
    doc.add_heading("5. Detailed Operation of Modules", level=1)

    modules = [
        ("Early Warning AI",
         "This module analyzes student attendance and internal marks to identify patterns that "
         "indicate academic difficulties. If a student consistently performs poorly or frequently "
         "misses classes, the system flags them as an at-risk student. Faculty members can then "
         "intervene early by offering academic support or counseling."),

        ("Seating Allocation System",
         "The seating allocation module automatically assigns students to examination halls and "
         "specific seat numbers. The system considers classroom capacity and student enrollment "
         "to distribute students efficiently. This prevents manual errors and ensures proper "
         "utilization of available space."),

        ("NLP Mind Map Tool",
         "Students can upload lecture notes or syllabus documents, and the system analyzes the "
         "content using natural language processing techniques. Key concepts are extracted and "
         "converted into visual mind maps that help students understand relationships between "
         "topics and improve their revision process."),

        ("Hall Ticket Generator",
         "This module generates personalized hall tickets once exam schedules and seating "
         "arrangements are finalized. Each hall ticket contains the student's name, register "
         "number, exam subjects, dates, room number, and seat number. Students can download "
         "the tickets directly from the platform."),

        ("Academic Calendar & Event Manager",
         "Administrators can publish academic events such as internal exams, semester start "
         "dates, holidays, and cultural programs. These events automatically appear on the "
         "dashboards of students and faculty members."
         )
    ]

    for mod, desc in modules:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{mod}: ")
        set_font(run, bold=True)
        run2 = p.add_run(desc)
        set_font(run2)

    # ------------------------------------------------
    # 6. SYSTEM WORKFLOW
    # ------------------------------------------------
    doc.add_heading("6. Integrated System Workflow", level=1)

    workflow_text = (
        "The workflow of VantageEdu begins with administrators entering student and course "
        "information into the system database. During the semester, attendance records and "
        "internal assessment data are continuously updated.\n\n"

        "When examination periods approach, administrators generate the exam schedule and "
        "activate the seating allocation module. Once seating arrangements are finalized, "
        "the hall ticket generator automatically creates downloadable hall tickets for each "
        "student.\n\n"

        "Students can log in to their dashboards to access their hall tickets, exam halls, "
        "and schedules. After the examinations are completed, all academic data remains "
        "stored within the system for reporting and future reference."
    )

    p = doc.add_paragraph()
    run = p.add_run(workflow_text)
    set_font(run)

    # ------------------------------------------------
    # SAVE FILE
    # ------------------------------------------------
    output_path = "/Users/rekha/.gemini/antigravity/brain/090f77a0-a0bf-488a-b422-56d169e63ad7/VantageEdu_Final_Official_Doc.docx"
    doc.save(output_path)

    print(f"\nDocumentation generated successfully!")
    print(f"File saved at: {output_path}")


if __name__ == "__main__":
    create_documentation()
